"""Completare Anexa 4 — Calcul pentru întreprinderile partenere sau legate.

Structura template (8 tabele, în ordinea din document):
  T0 — pag.1 "Calculul pentru tipurile..." (6×4): R2=solicitant, R3=partenere(A),
       R4=legate(B), R5=TOTAL. Coloane c1=salariați, c2=cifră, c3=active (mii).
  T1 — Secțiunea A, Tabelul A.1: lista partenerilor (11×7), R2..R9 date, R10=Total.
  T2 — FIȘA DE PARTENERIAT: date 100% partener (3×4, R2="Total").
  T3 — FIȘA DE PARTENERIAT: Tabelul A.2, rezultat proporțional (2×4, R1).
  T4 — Tabelul B1 (Cazul 1, consolidat) — NEUTILIZAT la Cazul 2.
  T5 — "Identificarea întreprinderilor incluse prin consolidare" (7×4) —
       RĂMÂNE GOL la Cazul 2 (firmele legate merg în Tabelul B2, nu aici).
  T6 — Tabelul B2: financiar legate (7×4): c0=denumire, c1..c3=financiare, R6=Total.
  T7 — FIȘA privind legătura: date legate (3×4, R2="Total").

Fișele (FIȘA DE PARTENERIAT P030.. și FIȘA privind legătura P106..) sunt blocuri
care se DUPLICĂ — câte una per întreprindere — prin deepcopy lxml pe elementele
<w:p>/<w:tbl> ale blocului.

Anul de referință apare DOAR în celule de tabel (T0/T2/T7 R0), de aceea înlocuirea
"2024"→an parcurge ȘI conținutul tabelelor (nu doar paragrafele de nivel-top).
"""
from __future__ import annotations

import argparse
import copy
import re
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("EROARE: python-docx lipsă. pip install python-docx", file=sys.stderr)
    sys.exit(2)


_PH = re.compile(r"[_\.…]{3,}")


def _set_cell(cell, text: str) -> None:
    """Setează textul unei celule fără a strica formatul (clear runs + add run)."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text))
    else:
        cell.add_paragraph(str(text))


def _fmt_sal(v) -> str:
    try:
        return f"{float(v):g}"
    except Exception:
        return str(v)


def _fmt_mii(v) -> str:
    try:
        return f"{float(v):.3f}"
    except Exception:
        return str(v)


def _fill_label(paragraph, value) -> bool:
    """Înlocuiește primul placeholder (___ / .... / ……) din paragraf cu ' value',
    păstrând eticheta. Curăță coada de placeholder din runurile următoare."""
    done = False
    for run in paragraph.runs:
        if not done:
            m = _PH.search(run.text)
            if m:
                run.text = run.text[: m.start()] + " " + str(value) + run.text[m.end():]
                done = True
                continue
        if done and run.text and set(run.text) <= set("_.… "):
            run.text = ""  # coadă de underscore-uri rămasă pe runurile următoare
    return done


def _replace_year(doc, an: int) -> None:
    """Înlocuiește 2024→an în TOATE paragrafele, inclusiv în celulele de tabel
    (unde stau de fapt antetele „Perioada de referință")."""
    old, new = "2024", str(an)
    if new == old:
        return

    def fix(p):
        if old not in p.text:
            return
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        # „2024" împărțit pe mai multe runuri → reconstruiește pe primul run
        if old in p.text and p.runs:
            merged = "".join(r.text for r in p.runs).replace(old, new)
            p.runs[0].text = merged
            for r in p.runs[1:]:
                r.text = ""

    for p in doc.paragraphs:
        fix(p)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    fix(p)


def _check_caz(doc, label_start: str) -> bool:
    """Bifează un „Cazul N" (checkbox = <w:sym Symbol F090>) → Wingdings F0FE."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(label_start):
            for sym in p._p.iter(qn("w:sym")):
                sym.set(qn("w:font"), "Wingdings")
                sym.set(qn("w:char"), "F0FE")
                return True
    return False


# ---------------- deepcopy bloc fișă ----------------

def _children(doc):
    return list(doc.element.body)


def _find_p_elem(doc, pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p._p
    return None


def _elem_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _block_range(doc, start_pred, end_pred):
    """(start_elem, end_elem). Dacă end_pred e None → end = ultimul w:p/w:tbl din
    body (fără sectPr)."""
    start = _find_p_elem(doc, start_pred)
    if start is None:
        return None, None
    kids = _children(doc)
    si = kids.index(start)
    if end_pred is None:
        end = None
        for e in kids[si:]:
            if e.tag in (qn("w:p"), qn("w:tbl")):
                end = e
        return start, end
    for e in kids[si + 1:]:
        if e.tag == qn("w:p") and end_pred(_elem_text(e)):
            return start, e
    return start, start


def _replicate(doc, start, end, n):
    """Listă de n instanțe (fiecare = listă de elemente). Prima = originalul;
    restul = deepcopy inserate imediat după bloc."""
    kids = _children(doc)
    si, ei = kids.index(start), kids.index(end)
    original = kids[si:ei + 1]
    instances = [original]
    anchor = end
    for _ in range(max(0, n - 1)):
        newb = [copy.deepcopy(e) for e in original]
        ref = anchor
        for ne in newb:
            ref.addnext(ne)
            ref = ne
        anchor = newb[-1]
        instances.append(newb)
    return instances


def _wrap(doc, elems):
    """Yield ('p', Paragraph) / ('t', Table) pentru elementele unui bloc."""
    for el in elems:
        if el.tag == qn("w:p"):
            yield "p", Paragraph(el, doc._body)
        elif el.tag == qn("w:tbl"):
            yield "t", Table(el, doc._body)


def _fill_parteneriat(doc, elems, p, an) -> None:
    """Completează un bloc FIȘA DE PARTENERIAT pentru un partener."""
    expect_proc = False
    for kind, obj in _wrap(doc, elems):
        if kind == "p":
            t = obj.text.strip()
            if t.startswith("Denumirea întreprinderii"):
                _fill_label(obj, p.get("denumire", ""))
            elif t.startswith("Adresa sediului social"):
                _fill_label(obj, p.get("adresa", ""))
            elif t.startswith("Codul unic"):
                _fill_label(obj, str(p.get("cui", "")))
            elif t.startswith("Numele, prenumele și funcția"):
                _fill_label(obj, p.get("reprezentant", ""))
            elif t.startswith("Indicați exact proporția deținută de întreprinderea solicitantă"):
                expect_proc = True
            elif expect_proc and _PH.search(t):
                _fill_label(obj, f"{p.get('procent', 0):.2f}%".replace(".", ","))
                expect_proc = False
        else:  # tabel
            rows = obj.rows
            if len(rows) >= 3:  # T2: 100% în R2 (Total)
                _set_cell(rows[2].cells[1], _fmt_sal(p.get("salariati", 0)))
                _set_cell(rows[2].cells[2], _fmt_mii(p.get("cifra_afaceri_mii_lei", 0)))
                _set_cell(rows[2].cells[3], _fmt_mii(p.get("active_totale_mii_lei", 0)))
            elif len(rows) == 2:  # T3 (A.2): proporțional în R1
                _set_cell(rows[1].cells[1], _fmt_sal(p.get("salariati_proportional", 0)))
                _set_cell(rows[1].cells[2], _fmt_mii(p.get("cifra_proportional", 0)))
                _set_cell(rows[1].cells[3], _fmt_mii(p.get("active_proportional", 0)))


def _fill_legatura(doc, elems, l, nr) -> None:
    """Completează un bloc FIȘA privind legătura pentru o firmă legată."""
    for kind, obj in _wrap(doc, elems):
        if kind == "p":
            t = obj.text.strip()
            if t.startswith("privind legătura"):
                _fill_label(obj, str(nr))
            elif t.startswith("Denumirea întreprinderii"):
                _fill_label(obj, l.get("denumire", ""))
            elif t.startswith("Adresa sediului social"):
                _fill_label(obj, l.get("adresa", ""))
            elif t.startswith("Codul unic"):
                _fill_label(obj, str(l.get("cui", "")))
            elif t.startswith("Numele, prenumele și funcția"):
                _fill_label(obj, l.get("reprezentant", ""))
        else:  # T7: R2 = Total
            rows = obj.rows
            if len(rows) >= 3:
                _set_cell(rows[2].cells[1], _fmt_sal(l.get("salariati", 0)))
                _set_cell(rows[2].cells[2], _fmt_mii(l.get("cifra_afaceri_mii_lei", 0)))
                _set_cell(rows[2].cells[3], _fmt_mii(l.get("active_totale_mii_lei", 0)))


def fill_anexa4(
    template_path: Path,
    out_path: Path,
    solicitant: dict,
    parteneri: list[dict],
    legate: list[dict],
    an_referinta: int,
    caz: int = 2,
) -> Path:
    """
    solicitant: {salariati, cifra_afaceri_mii_lei, active_totale_mii_lei}
                (datele PROPRII ale solicitantului — rândul 1 din tabelul pag.1)
    parteneri:  [{denumire, adresa, cui, reprezentant, procent,
                  salariati, cifra_afaceri_mii_lei, active_totale_mii_lei,
                  salariati_proportional, cifra_proportional, active_proportional}]
    legate:     [{denumire, adresa, cui, reprezentant,
                  salariati, cifra_afaceri_mii_lei, active_totale_mii_lei}]
    caz:        1 sau 2 (Secțiunea B) — pentru ROMCARBON & filiale individuale = 2.
    """
    if len(parteneri) > 8:
        raise ValueError(f"Anexa 4 — max 8 parteneri; primit {len(parteneri)}")
    if len(legate) > 5:
        raise ValueError(f"Anexa 4 — max 5 legate (Tabel B2); primit {len(legate)}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, out_path)
    doc = Document(out_path)
    if len(doc.tables) < 8:
        raise RuntimeError(f"Anexa 4 — așteptat 8 tabele, găsite {len(doc.tables)}")

    # ---- agregate ----
    A_sal = sum(p.get("salariati_proportional", 0) for p in parteneri)
    A_ca = sum(p.get("cifra_proportional", 0) for p in parteneri)
    A_act = sum(p.get("active_proportional", 0) for p in parteneri)
    B_sal = sum(l.get("salariati", 0) for l in legate)
    B_ca = sum(l.get("cifra_afaceri_mii_lei", 0) for l in legate)
    B_act = sum(l.get("active_totale_mii_lei", 0) for l in legate)
    s_sal = solicitant.get("salariati", 0)
    s_ca = solicitant.get("cifra_afaceri_mii_lei", 0)
    s_act = solicitant.get("active_totale_mii_lei", 0)

    # ---------------- T0 — pag.1: solicitant / A / B / TOTAL ----------------
    t0 = doc.tables[0]
    t0_rows = [
        (2, s_sal, s_ca, s_act),
        (3, A_sal, A_ca, A_act),
        (4, B_sal, B_ca, B_act),
        (5, s_sal + A_sal + B_sal, s_ca + A_ca + B_ca, s_act + A_act + B_act),
    ]
    for ri, sal, ca, act in t0_rows:
        if ri < len(t0.rows):
            cells = t0.rows[ri].cells
            _set_cell(cells[1], _fmt_sal(sal))
            _set_cell(cells[2], _fmt_mii(ca))
            _set_cell(cells[3], _fmt_mii(act))

    # ---------------- T1 — Secțiunea A: lista partenerilor ----------------
    t1 = doc.tables[1]
    for i, p in enumerate(parteneri):
        ri = 2 + i
        if ri >= len(t1.rows):
            break
        c = t1.rows[ri].cells
        _set_cell(c[0], f"{i + 1}. {p.get('denumire', '')}")
        _set_cell(c[1], str(p.get("adresa", "")))
        _set_cell(c[2], str(p.get("cui", "")))
        _set_cell(c[3], str(p.get("reprezentant", "")))
        _set_cell(c[4], _fmt_sal(p.get("salariati_proportional", 0)))
        _set_cell(c[5], _fmt_mii(p.get("cifra_proportional", 0)))
        _set_cell(c[6], _fmt_mii(p.get("active_proportional", 0)))
    if len(t1.rows) >= 11:
        tr = t1.rows[10].cells
        _set_cell(tr[4], _fmt_sal(A_sal))
        _set_cell(tr[5], _fmt_mii(A_ca))
        _set_cell(tr[6], _fmt_mii(A_act))

    # ---------------- T5 — RĂMÂNE GOL (Cazul 2: legatele merg în B2) -------
    # (nu se completează nimic — corectură feedback)

    # ---------------- T6 — Tabelul B2: denumire + financiare legate -------
    t6 = doc.tables[6]
    for i, l in enumerate(legate):
        ri = 1 + i
        if ri >= len(t6.rows):
            break
        c = t6.rows[ri].cells
        _set_cell(c[0], f"{i + 1}. {l.get('denumire', '')} *)")
        _set_cell(c[1], _fmt_sal(l.get("salariati", 0)))
        _set_cell(c[2], _fmt_mii(l.get("cifra_afaceri_mii_lei", 0)))
        _set_cell(c[3], _fmt_mii(l.get("active_totale_mii_lei", 0)))
    if len(t6.rows) >= 7:
        tr = t6.rows[6].cells
        _set_cell(tr[1], _fmt_sal(B_sal))
        _set_cell(tr[2], _fmt_mii(B_ca))
        _set_cell(tr[3], _fmt_mii(B_act))

    # ---------------- Secțiunea B: bifează Cazul corect -------------------
    _check_caz(doc, f"Cazul {caz}:")

    # ---------------- FIȘA DE PARTENERIAT — una per partener --------------
    if parteneri:
        ps, pe = _block_range(
            doc,
            lambda t: t.startswith("FIȘA DE PARTENERIAT"),
            lambda t: t.startswith("Aceste date se vor introduce în Tabelul A.1"),
        )
        if ps is not None and pe is not None:
            for inst, p in zip(_replicate(doc, ps, pe, len(parteneri)), parteneri):
                _fill_parteneriat(doc, inst, p, an_referinta)

    # ---------------- FIȘA privind legătura — una per firmă legată --------
    if legate:
        ls, le = _block_range(doc, lambda t: t == "FIȘA", None)
        if ls is not None and le is not None:
            for i, (inst, l) in enumerate(zip(_replicate(doc, ls, le, len(legate)), legate)):
                _fill_legatura(doc, inst, l, nr=i + 1)

    # ---------------- Anul de referință (2024→an) în tot documentul -------
    _replace_year(doc, an_referinta)

    doc.save(out_path)
    return out_path


# -------------------- CLI smoke test --------------------

def _smoke():
    sys.path.insert(0, str(Path(__file__).parent))
    from _paths import templates_dir
    tpl = templates_dir() / "Anexa-4-Calcul-intreprinderi.docx"
    out = Path.cwd() / f"_smoke_Anexa4_{date.today().isoformat()}.docx"
    solicitant = {"salariati": 604, "cifra_afaceri_mii_lei": 196460.720, "active_totale_mii_lei": 239172.373}
    parteneri = [{
        "denumire": "YENKI SRL", "adresa": "Buzău", "cui": "22404794", "reprezentant": "-",
        "procent": 33.34, "salariati": 0, "cifra_afaceri_mii_lei": 12.605, "active_totale_mii_lei": 106.063,
        "salariati_proportional": 0.0, "cifra_proportional": 4.203, "active_proportional": 35.361,
    }]
    legate = [
        {"denumire": "RC ENERGO INSTALL SRL", "adresa": "Buzău", "cui": "17701071", "reprezentant": "-",
         "salariati": 52, "cifra_afaceri_mii_lei": 13993.962, "active_totale_mii_lei": 14735.124},
        {"denumire": "LIVINGJUMBO INDUSTRY SA", "adresa": "Buzău", "cui": "14419107", "reprezentant": "-",
         "salariati": 420, "cifra_afaceri_mii_lei": 93304.833, "active_totale_mii_lei": 42600.535},
    ]
    result = fill_anexa4(tpl, out, solicitant, parteneri, legate, an_referinta=2025, caz=2)
    print(f"Smoke OK: {result}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--smoke", action="store_true")
    args = ap.parse_args()
    if args.smoke:
        _smoke()
