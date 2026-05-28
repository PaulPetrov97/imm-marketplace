"""Generator sinteză — produce MD afișat inline + DOCX echivalent.

Output structurat: profil firmă → matrice asociați → clasificare consolidată →
verdict 'Există legături: DA/NU/NECLAR' cu raționament.

Cap lungime: 1500 cuvinte (≈1-3 pagini A4).
"""
from __future__ import annotations

import argparse
import sys
from datetime import date, datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("EROARE: python-docx lipsă. pip install python-docx", file=sys.stderr)
    sys.exit(2)


# -------------------- MD generator --------------------

def _md_table(headers: list[str], rows: list[list[str]]) -> str:
    if not rows:
        rows = [["—"] * len(headers)]
    out = ["| " + " | ".join(headers) + " |",
           "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        cells = [str(c).replace("|", "\\|") for c in row]
        # Pad / truncate la len(headers)
        cells = (cells + [""] * len(headers))[:len(headers)]
        out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def generate_md(
    solicitant: dict,
    asociati_matrix: list[dict],
    clasificare: dict,
    parteneri: list[dict],
    legate: list[dict],
    neclar: list[dict],
    agregat: dict,
    an_referinta: int,
    audit_sources: list[str],
) -> str:
    """Generează sinteza ca string Markdown."""
    today = date.today().strftime("%d %B %Y")
    verdict = _verdict_label(clasificare, neclar)

    md = [
        f"# Sinteză preliminară IMM — {solicitant.get('denumire', '(necunoscut)')} "
        f"(CUI {solicitant.get('cui', '?')})",
        "",
        f"**Data analizei:** {today} &nbsp;·&nbsp; "
        f"**An referință financiar:** {an_referinta} &nbsp;·&nbsp; "
        f"**Surse:** {', '.join(audit_sources) or 'termene.ro'}",
        "",
        "## 1. Profil întreprindere solicitantă",
        "",
        _md_table(
            ["Câmp", "Valoare"],
            [
                ["Denumire", solicitant.get("denumire", "—")],
                ["CUI", solicitant.get("cui", "—")],
                ["Adresă sediu social", solicitant.get("adresa", "—")],
                ["CAEN principal",
                 f"{solicitant.get('caen_principal', {}).get('cod', '—')} — "
                 f"{solicitant.get('caen_principal', {}).get('label', '')}"],
                ["Reprezentant legal",
                 ", ".join(f"{r.get('nume','?')} ({r.get('functie','?')})"
                           for r in solicitant.get("reprezentant_legal", []))
                 or "—"],
                ["Stare", solicitant.get("stare", "—")],
                ["Salariați (an N)", f"{solicitant.get('salariati', '—')}"],
                ["Cifră afaceri (an N)",
                 f"{solicitant.get('cifra_afaceri_lei', 0):,.2f} lei"
                 if solicitant.get('cifra_afaceri_lei') else "—"],
                ["Active totale (an N)",
                 f"{solicitant.get('active_totale_lei', 0):,.2f} lei"
                 if solicitant.get('active_totale_lei') else "—"],
            ]
        ),
        "",
        "## 2. Matrice asociați și suprapuneri",
        "",
    ]

    if asociati_matrix:
        rows = []
        for a in asociati_matrix:
            alte_str = "; ".join(
                f"{af.get('denumire','?')} ({af.get('procent_in','?')}%, "
                f"CAEN {af.get('caen','?')})"
                for af in a.get("alte_firme", [])
            ) or "—"
            rows.append([
                a.get("nume", "?"),
                a.get("tip", "?"),
                f"{a.get('procent', 0):.2f}%",
                alte_str[:120] + ("…" if len(alte_str) > 120 else ""),
                a.get("piata_invecinata", "—"),
            ])
        md.append(_md_table(
            ["Asociat", "Tip", "% deținere", "Alte firme ≥25%", "Piață învecinată?"],
            rows
        ))
    else:
        md.append("_Niciun asociat extras._")
    md.append("")

    md.extend([
        "## 3. Parteneri identificați",
        "",
    ])
    if parteneri:
        rows = [[p.get("nume", "?"), p.get("cui", "?"), f"{p.get('procent', 0):.2f}%",
                 p.get("sens", "—"), p.get("rol", "—")] for p in parteneri]
        md.append(_md_table(["Denumire", "CUI", "Procent", "Sens", "Rol"], rows))
    else:
        md.append("_Niciun partener identificat._")
    md.append("")

    md.extend([
        "## 4. Întreprinderi legate identificate",
        "",
    ])
    if legate:
        rows = [[l.get("nume", "?"), l.get("cui", "?"), f"{l.get('procent', 0):.2f}%",
                 l.get("motiv", "—")] for l in legate]
        md.append(_md_table(["Denumire", "CUI", "Procent", "Motiv"], rows))
    else:
        md.append("_Nicio întreprindere legată identificată._")
    md.append("")

    if neclar:
        md.extend([
            "## 5. Cazuri neclare",
            "",
            _md_table(
                ["Subiect", "Motiv"],
                [[str(n.get("subiect", "?")), str(n.get("motiv", "?"))] for n in neclar]
            ),
            "",
        ])

    md.extend([
        "## 6. Clasificare consolidată",
        "",
        _md_table(
            ["Criteriu", "Solicitant", "+ Parteneri (proporțional)",
             "+ Legate (integral)", "TOTAL", "Prag IMM", "Conformitate"],
            _build_classification_rows(solicitant, parteneri, legate, agregat)
        ),
        "",
        f"**Tip relațional:** **{clasificare.get('tip', '—')}**",
        f"**Categorie de mărime:** **{clasificare.get('marime', '—')}**",
        "",
        f"## 7. Există legături? **{verdict}**",
        "",
        clasificare.get("reason", ""),
        "",
    ])

    if neclar:
        md.extend([
            "### Informații lipsă pentru clarificare definitivă",
            "",
        ])
        for n in neclar:
            md.append(f"- **{n.get('subiect','?')}** — {n.get('motiv','?')}")
        md.append("")

    md.extend([
        "---",
        "",
        "_Sinteză generată automat de pluginul `imm-analiza-ro` (Claude Code). "
        "Decision engine bazat pe Legea 346/2004 și Recomandarea CE 2003/361/EC. "
        "A se valida cu declarațiile pe propria răspundere ale solicitantului._",
    ])
    return "\n".join(md)


def _verdict_label(clasificare: dict, neclar: list[dict]) -> str:
    tip = clasificare.get("tip", "")
    if tip in {"PARTENERA", "LEGATA"}:
        if neclar:
            return "DA (cu cazuri NECLARE de validat)"
        return "DA"
    if tip == "AUTONOMA":
        if neclar:
            return "NECLAR"
        return "NU"
    if tip == "NU_IMM":
        return "NU SE ÎNCADREAZĂ CA IMM"
    return "NECLAR"


def _build_classification_rows(
    solicitant: dict, parteneri: list[dict], legate: list[dict], agregat: dict,
) -> list[list[str]]:
    sol_sal = solicitant.get("salariati", 0)
    sol_ca = solicitant.get("cifra_afaceri_lei", 0)
    sol_act = solicitant.get("active_totale_lei", 0)
    p_sal = sum(p.get("salariati_proportional", 0) for p in parteneri)
    p_ca = sum(p.get("cifra_proportional_lei", 0) for p in parteneri)
    p_act = sum(p.get("active_proportional_lei", 0) for p in parteneri)
    l_sal = sum(l.get("salariati", 0) for l in legate)
    l_ca = sum(l.get("cifra_afaceri_lei", 0) for l in legate)
    l_act = sum(l.get("active_totale_lei", 0) for l in legate)
    tot_sal = agregat.get("salariati", sol_sal + p_sal + l_sal)
    tot_ca = agregat.get("cifra_afaceri_lei", sol_ca + p_ca + l_ca)
    tot_act = agregat.get("active_totale_lei", sol_act + p_act + l_act)
    tot_ca_eur = agregat.get("cifra_afaceri_eur", 0)
    tot_act_eur = agregat.get("active_totale_eur", 0)
    return [
        ["Salariați", f"{sol_sal:.0f}", f"{p_sal:.2f}", f"{l_sal:.0f}",
         f"{tot_sal:.2f}", "<250", "✓" if tot_sal < 250 else "✗"],
        ["Cifră afaceri (lei)", f"{sol_ca:,.0f}", f"{p_ca:,.0f}", f"{l_ca:,.0f}",
         f"{tot_ca:,.0f}", "≤€50M",
         "✓" if tot_ca_eur <= 50_000_000 else "✗"],
        ["Active totale (lei)", f"{sol_act:,.0f}", f"{p_act:,.0f}", f"{l_act:,.0f}",
         f"{tot_act:,.0f}", "≤€43M",
         "✓" if tot_act_eur <= 43_000_000 else "✗"],
    ]


# -------------------- DOCX generator --------------------

def generate_docx(md_text: str, out_path: Path) -> Path:
    """Convertor simplu MD → DOCX (titluri, paragrafe, tabele)."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        # Heading
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        # Table block
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            # Parse
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("| ").split("|")]
                rows.append(cells)
            if len(rows) >= 2 and all(c.strip("-: ") == "" for c in rows[1]):
                rows.pop(1)  # remove separator row
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        if ci < len(tbl.rows[ri].cells):
                            tbl.rows[ri].cells[ci].text = val
        # Bullet
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        # Horizontal rule
        elif line.startswith("---"):
            doc.add_paragraph("_" * 40)
        elif line.strip():
            # Strip simple markdown bold markers for docx text
            p = doc.add_paragraph()
            txt = line.replace("**", "").replace("`", "")
            p.add_run(txt)
        else:
            doc.add_paragraph("")
        i += 1

    doc.save(out_path)
    return out_path


# -------------------- High-level API --------------------

def generate_sinteza(
    solicitant: dict,
    asociati_matrix: list[dict],
    clasificare: dict,
    parteneri: list[dict],
    legate: list[dict],
    neclar: list[dict],
    agregat: dict,
    an_referinta: int,
    audit_sources: list[str],
    out_dir: Path,
) -> tuple[Path, Path]:
    """Produce ambele fișiere și returnează path-urile (md, docx)."""
    md_text = generate_md(solicitant, asociati_matrix, clasificare, parteneri,
                          legate, neclar, agregat, an_referinta, audit_sources)
    cui = solicitant.get("cui", "UNKNOWN")
    md_path = out_dir / f"02_Sinteza_{cui}.md"
    docx_path = out_dir / f"02_Sinteza_{cui}.docx"
    md_path.write_text(md_text, encoding="utf-8")
    generate_docx(md_text, docx_path)
    return md_path, docx_path


# -------------------- Smoke test --------------------

def _smoke():
    solicitant = {
        "denumire": "ROMACTIV BUSINESS CONSULTING SRL",
        "cui": "RO14186770",
        "adresa": "Str. Constantin Tănase 12, București, Sector 2",
        "caen_principal": {"cod": "7022", "label": "Activități de consultanță în management"},
        "reprezentant_legal": [{"nume": "POPESCU PAUL", "functie": "Administrator"}],
        "stare": "Funcțională",
        "salariati": 8,
        "cifra_afaceri_lei": 1_500_000.0,
        "active_totale_lei": 850_000.0,
    }
    asociati = [
        {"nume": "POPESCU PAUL", "tip": "PF", "procent": 100.0,
         "alte_firme": [], "piata_invecinata": "—"},
    ]
    clasificare = {"tip": "AUTONOMA", "marime": "micro", "reason": "Asociat unic PF, fără alte dețineri."}
    agregat = {
        "salariati": 8, "cifra_afaceri_lei": 1_500_000.0, "active_totale_lei": 850_000.0,
        "cifra_afaceri_eur": 301_456.0, "active_totale_eur": 170_825.0, "curs_eur": 4.9759,
    }
    out = Path.cwd() / "_smoke_sinteza"
    out.mkdir(exist_ok=True)
    md_path, docx_path = generate_sinteza(
        solicitant, asociati, clasificare, [], [], [], agregat,
        an_referinta=2024, audit_sources=["termene.ro (premium)"],
        out_dir=out,
    )
    print(f"MD : {md_path} ({md_path.stat().st_size} bytes)")
    print(f"DOCX: {docx_path} ({docx_path.stat().st_size} bytes)")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--smoke", action="store_true")
    args = ap.parse_args()
    if args.smoke:
        _smoke()
