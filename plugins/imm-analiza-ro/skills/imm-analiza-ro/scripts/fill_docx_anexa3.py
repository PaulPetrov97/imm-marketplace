"""Completare Anexa 3 — Declarație privind încadrarea întreprinderii în categoria IMM.

Pași:
1. shutil.copy(template, out) — copie nouă.
2. Înlocuiește placeholder-urile (underscore / puncte / …) păstrând ÎNTOTDEAUNA
   eticheta — ex. cuvântul "Numele" NU se șterge, valoarea se scrie LÂNGĂ el.
3. Bifează una SAU MAI MULTE căsuțe (autonomă / parteneră / legată) — o firmă
   poate fi simultan parteneră ȘI legată — prin înlocuirea Wingdings unchecked
   cu CHECKED (U+F0FE), forțând fontul Wingdings pe runul bifei.
4. Completează tabelul financiar (Tabel 2).
5. Completează blocul de semnătură: Numele = persoana care întocmește anexa
   (semnatarul — se întreabă la începutul task-ului), Funcția, Data semnării.
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("EROARE: python-docx lipsă. Rulează: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# Caractere Wingdings pentru checkbox (private use area)
# Template-ul folosește un caracter unchecked pe care îl detectăm dinamic la rulare;
# înlocuim cu U+2611 (BALLOT BOX WITH CHECK) sau U+F0FE (Wingdings) după font.
UNCHECKED_CANDIDATES = ["", "", "☐", "☐"]
CHECKED_WINGDINGS = ""
CHECKED_UNICODE = "☑"


TIP_TO_ROW_INDEX = {
    "AUTONOMA": 0,
    "PARTENERA": 1,
    "LEGATA": 2,
}


# Placeholder = secvență de ≥3 caractere de tip underscore / punct / elipsă (…)
_PLACEHOLDER_RE = re.compile(r"[_\.…]{3,}")


def _fill_labeled(paragraph, label: str, value: str) -> bool:
    """Completează un paragraf-etichetă fără a șterge eticheta.

    Caută în paragraful care începe cu `label` primul run care conține o
    secvență placeholder (___ / ..... / ……) și înlocuiește DOAR acea secvență
    cu valoarea — textul etichetei rămâne intact chiar dacă etichetă +
    placeholder stau în același run (ex. "Numele......").
    """
    if not paragraph.text.startswith(label):
        return False
    for run in paragraph.runs:
        m = _PLACEHOLDER_RE.search(run.text)
        if m:
            run.text = run.text[:m.start()] + " " + value
            return True
    return False


def _tick_checkbox(cell) -> bool:
    """Bifează căsuța dintr-o celulă, forțând randarea Wingdings.

    Template-ul ține caracterul de checkbox într-un run de text FĂRĂ font setat
    (run.font.name = None), deci varianta veche scria ☑ (U+2611) în Times — bifa
    ieșea inconsistentă cu căsuțele goale. Aici înlocuim cu CHECKED_WINGDINGS
    (U+F0FE) ȘI setăm fontul runului pe Wingdings ca să randeze corect.
    """
    for para in cell.paragraphs:
        for run in para.runs:
            for unchecked in UNCHECKED_CANDIDATES:
                if unchecked and unchecked in run.text:
                    run.text = run.text.replace(unchecked, CHECKED_WINGDINGS, 1)
                    run.font.name = "Wingdings"
                    return True
    return False


def fill_anexa3(
    template_path: Path,
    out_path: Path,
    solicitant: dict,
    totals_ref: dict,
    totals_prev: dict | None,
    tip,  # str SAU listă/set: "LEGATA", sau {"PARTENERA","LEGATA"} când e și parteneră și legată
    modif_categorie: bool = False,
    semnatar_nume: str | None = None,
    semnatar_functie: str | None = None,
) -> Path:
    """
    solicitant: {denumire, adresa, cui, reprezentant_nume, reprezentant_functie}
    totals_ref: {salariati, cifra_afaceri_mii_lei, active_totale_mii_lei, an}
    totals_prev: idem, pentru exercițiul anterior (opțional)
    tip: "AUTONOMA" | "PARTENERA" | "LEGATA"
    modif_categorie: True dacă datele financiare au modificat categoria față de N-1
    semnatar_nume / semnatar_functie: persoana care întocmește/semnează anexa
        (blocul "Numele... / Funcția:..." de la final). Dacă lipsesc, se
        folosește reprezentantul legal. Eticheta "Numele" NU se șterge —
        valoarea se scrie lângă ea.
    """
    tips = [tip] if isinstance(tip, str) else list(tip)
    target_idx = set()
    for _t in tips:
        if _t not in TIP_TO_ROW_INDEX:
            raise ValueError(f"Tip invalid: {_t}. Trebuie AUTONOMA/PARTENERA/LEGATA")
        target_idx.add(TIP_TO_ROW_INDEX[_t])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, out_path)
    doc = Document(out_path)

    # ---- 1. Identificare în paragrafe ----
    placeholders = {
        "Denumirea întreprinderii": solicitant.get("denumire", ""),
        "Adresa sediului social":   solicitant.get("adresa", ""),
        "Cod unic de înregistrare": solicitant.get("cui", ""),
        "Numele și funcția":        f"{solicitant.get('reprezentant_nume','')}, "
                                    f"{solicitant.get('reprezentant_functie','')}",
    }
    for p in doc.paragraphs:
        for label, value in placeholders.items():
            # sare peste valorile goale (ex. reprezentant neconfirmat) — placeholder-ul rămâne
            if value and value.strip(" ,") and _fill_labeled(p, label, value):
                break

    # ---- 2. Tabel 0: bifa pe rândul corect ----
    if len(doc.tables) < 1:
        raise RuntimeError("Anexa 3 — tabelul de tip întreprindere lipsește")
    t0 = doc.tables[0]
    if len(t0.rows) < 3:
        raise RuntimeError(f"Anexa 3 — tabel tip are {len(t0.rows)} rânduri, aștept 3")
    for ri in range(3):
        cell = t0.rows[ri].cells[0]
        if ri in target_idx:
            ok = _tick_checkbox(cell)
            if not ok:
                # Fallback: scrie un X (vizibil chiar dacă font-ul nu randează)
                cell.paragraphs[0].add_run(" X").bold = True

    # ---- 3. Tabel 1: date financiare ----
    if len(doc.tables) < 2:
        raise RuntimeError("Anexa 3 — tabelul financiar lipsește")
    t1 = doc.tables[1]
    if len(t1.rows) < 3:
        raise RuntimeError(f"Anexa 3 — tabel financiar are {len(t1.rows)} rânduri")

    # Antet an de referință (row 0)
    if "an" in totals_ref:
        t1.rows[0].cells[0].text = f"Exercițiul financiar de referință — anul {totals_ref['an']}"

    # Row 2 = primul exercițiu (referință)
    _set_cell(t1.rows[2].cells[0], f"{totals_ref.get('salariati', 0):g}")
    _set_cell(t1.rows[2].cells[1], f"{totals_ref.get('cifra_afaceri_mii_lei', 0):.3f}")
    _set_cell(t1.rows[2].cells[2], f"{totals_ref.get('active_totale_mii_lei', 0):.3f}")

    # Tabel III are UN SINGUR rând de date (R2 = anul de referință). Rândul R3
    # rămâne GOL — datele exercițiului anterior NU se trec aici (feedback Paul).
    _ = totals_prev  # păstrat în semnătură pentru compatibilitate, dar neutilizat aici

    # ---- 4. Tabel 2: modificare categorie (Da/Nu) ----
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        if len(t2.rows) >= 1 and len(t2.rows[0].cells) >= 2:
            # Adaugă marcajul în celula 2
            mark = "X DA" if modif_categorie else "X NU"
            t2.rows[0].cells[1].paragraphs[0].add_run(f"\n[{mark}]").bold = True

    # ---- 5. Bloc semnătură: Numele / Funcția / Data semnării ----
    # Eticheta rămâne MEREU; valoarea se scrie lângă ea.
    # Semnatarul se completează DOAR dacă e confirmat. Dacă lipsește, eticheta
    # „Numele"/„Funcția" rămâne cu placeholder-ul punctat pentru completare manuală
    # (firma poate avea mai mulți administratori — semnatarul îl alege clientul).
    semn_nume = (semnatar_nume or "").strip()
    semn_functie = (semnatar_functie or "").strip()
    today_str = date.today().strftime("%d.%m.%Y")
    for p in doc.paragraphs:
        txt = p.text
        if txt.startswith("Numele") and not txt.startswith("Numele și funcția"):
            if semn_nume:
                _fill_labeled(p, "Numele", semn_nume)
        elif txt.startswith("Funcția") or txt.startswith("Functia"):
            if semn_functie:
                _fill_labeled(p, "Funcția" if txt.startswith("Funcția") else "Functia",
                              semn_functie)
        elif "Data semnării" in txt:
            if not _fill_labeled(p, "Data semnării", today_str):
                for run in p.runs:
                    if "……" in run.text or "..." in run.text:
                        run.text = re.sub(r"[…\.]{2,}", " " + today_str, run.text,
                                          count=1)
                        break

    doc.save(out_path)
    return out_path


def _set_cell(cell, text: str) -> None:
    """Setează textul unei celule fără a strica formatarea — clear + add run."""
    # Curăță paragrafele existente
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    # Pune textul în primul paragraf (sau adaugă unul)
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


# -------------------- CLI smoke test --------------------

def _smoke():
    sys.path.insert(0, str(Path(__file__).parent))
    from _paths import templates_dir
    tpl = templates_dir() / "Anexa-3-Declaratie-incadrare-IMM.docx"
    out = Path.cwd() / f"_smoke_Anexa3_{date.today().isoformat()}.docx"
    solicitant = {
        "denumire": "ROMACTIV BUSINESS CONSULTING SRL",
        "adresa": "Str. Constantin Tănase nr. 12, București, Sector 2",
        "cui": "RO14186770",
        "reprezentant_nume": "POPESCU PAUL",
        "reprezentant_functie": "Administrator",
    }
    totals_ref = {
        "an": 2024,
        "salariati": 8,
        "cifra_afaceri_mii_lei": 1500.000,
        "active_totale_mii_lei": 850.000,
    }
    totals_prev = {
        "salariati": 7,
        "cifra_afaceri_mii_lei": 1200.000,
        "active_totale_mii_lei": 700.000,
    }
    result = fill_anexa3(tpl, out, solicitant, totals_ref, totals_prev, tip="PARTENERA",
                        modif_categorie=False,
                        semnatar_nume="IONESCU MARIA",
                        semnatar_functie="Consultant fonduri europene")
    print(f"Smoke OK: {result}")

    # Verificare: eticheta "Numele" pastrata + semnatarul scris langa ea
    chk = Document(result)
    ok_nume = ok_func = False
    for p in chk.paragraphs:
        t = p.text
        if t.startswith("Numele") and "și funcția" not in t and "IONESCU MARIA" in t:
            ok_nume = True
            print(f"  [OK] semnatura nume: {t!r}")
        if (t.startswith("Funcția") or t.startswith("Functia")) and "Consultant" in t:
            ok_func = True
            print(f"  [OK] semnatura functie: {t!r}")
    assert ok_nume, "Eticheta 'Numele' + semnatar NU au fost completate corect!"
    assert ok_func, "Eticheta 'Funcția' + functie semnatar NU au fost completate corect!"
    print("  Etichetele pastrate, semnatarul completat langa ele — OK")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--smoke", action="store_true")
    args = ap.parse_args()
    if args.smoke:
        _smoke()
