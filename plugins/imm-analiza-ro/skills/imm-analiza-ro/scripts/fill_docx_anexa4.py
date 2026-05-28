"""Completare Anexa 4 — Calcul pentru întreprinderile partenere sau legate.

Tabele așteptate (per analiza structurală):
  T0 — sumar (header "Perioada de referință")
  T1 — Secțiunea A: lista partenerilor (11×7)
  T2 — Secțiunea A: agregat parteneri (3×4)
  T3 — Tabel "Procent" (2×4)
  T4 — Tabel financiar per partener (2×4)
  T5 — Secțiunea B: identificare legate (7×4)
  T6 — Secțiunea B: financiar legate (7×4)
  T7 — Secțiunea B: agregat (3×4)

Completăm rândurile de date din T1, T5, T6 + totalurile.
"""
from __future__ import annotations

import argparse
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("EROARE: python-docx lipsă. pip install python-docx", file=sys.stderr)
    sys.exit(2)


def _set_cell(cell, text: str) -> None:
    """Setează textul unei celule fără a strica formatul."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


def fill_anexa4(
    template_path: Path,
    out_path: Path,
    parteneri: list[dict],
    legate: list[dict],
    agregat: dict,
    an_referinta: int,
) -> Path:
    """
    parteneri item: {denumire, adresa, cui, reprezentant, procent (0–100),
                     salariati, cifra_afaceri_mii_lei, active_totale_mii_lei,
                     salariati_proportional, cifra_proportional, active_proportional}
    legate item: {denumire, adresa, cui, reprezentant,
                  salariati, cifra_afaceri_mii_lei, active_totale_mii_lei}
    agregat: {A_salariati_sum, A_cifra_sum, A_active_sum,
              B_salariati_sum, B_cifra_sum, B_active_sum}
    """
    if len(parteneri) > 8:
        raise ValueError(f"Anexa 4 — max 8 parteneri în T1; primit {len(parteneri)}")
    if len(legate) > 5:
        raise ValueError(f"Anexa 4 — max 5 legate în T5/T6; primit {len(legate)}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, out_path)
    doc = Document(out_path)

    if len(doc.tables) < 7:
        raise RuntimeError(
            f"Anexa 4 — așteptat ≥7 tabele, găsite {len(doc.tables)}"
        )

    # ---------------- T1 — Secțiunea A: lista partenerilor ----------------
    t1 = doc.tables[1]
    # Rândurile de date sunt 2..9 (8 sloturi), antet în rândurile 0–1, total în 10
    for i, p in enumerate(parteneri):
        ri = 2 + i
        if ri >= len(t1.rows):
            break
        row = t1.rows[ri].cells
        _set_cell(row[0], str(p.get("denumire", "")))
        _set_cell(row[1], str(p.get("adresa", "")))
        _set_cell(row[2], str(p.get("cui", "")))
        _set_cell(row[3], str(p.get("reprezentant", "")))
        # Coloane 4–6: salariați/cifră/active PROPORȚIONAL (= % × valoare)
        _set_cell(row[4], f"{p.get('salariati_proportional', 0):.2f}")
        _set_cell(row[5], f"{p.get('cifra_proportional', 0):.3f}")
        _set_cell(row[6], f"{p.get('active_proportional', 0):.3f}")

    # Rândul total — index 10 dacă există 11 rânduri
    if len(t1.rows) >= 11:
        total_row = t1.rows[10].cells
        sum_sal = sum(p.get("salariati_proportional", 0) for p in parteneri)
        sum_ca = sum(p.get("cifra_proportional", 0) for p in parteneri)
        sum_act = sum(p.get("active_proportional", 0) for p in parteneri)
        _set_cell(total_row[4], f"{sum_sal:.2f}")
        _set_cell(total_row[5], f"{sum_ca:.3f}")
        _set_cell(total_row[6], f"{sum_act:.3f}")

    # ---------------- T5 — Secțiunea B: identificare legate ----------------
    t5 = doc.tables[5]
    for i, l in enumerate(legate):
        ri = 2 + i
        if ri >= len(t5.rows):
            break
        cells = t5.rows[ri].cells
        _set_cell(cells[0], str(l.get("denumire", "")))
        _set_cell(cells[1], str(l.get("adresa", "")))
        _set_cell(cells[2], str(l.get("cui", "")))
        _set_cell(cells[3], str(l.get("reprezentant", "")))

    # ---------------- T6 — Secțiunea B: financiar legate ----------------
    t6 = doc.tables[6]
    # Rândurile 1..5 sunt date (5 sloturi), rând 6 = Total
    for i, l in enumerate(legate):
        ri = 1 + i
        if ri >= len(t6.rows):
            break
        cells = t6.rows[ri].cells
        _set_cell(cells[1], f"{l.get('salariati', 0):.2f}")
        _set_cell(cells[2], f"{l.get('cifra_afaceri_mii_lei', 0):.3f}")
        _set_cell(cells[3], f"{l.get('active_totale_mii_lei', 0):.3f}")

    if len(t6.rows) >= 7:
        tot_b = t6.rows[6].cells
        sum_sal_b = sum(l.get("salariati", 0) for l in legate)
        sum_ca_b = sum(l.get("cifra_afaceri_mii_lei", 0) for l in legate)
        sum_act_b = sum(l.get("active_totale_mii_lei", 0) for l in legate)
        _set_cell(tot_b[1], f"{sum_sal_b:.2f}")
        _set_cell(tot_b[2], f"{sum_ca_b:.3f}")
        _set_cell(tot_b[3], f"{sum_act_b:.3f}")

    # ---------------- Update anul în antetele tabelelor (înlocuire textuală) ----------------
    for p in doc.paragraphs:
        if "Perioada de referință" in p.text and "2024" in p.text and an_referinta != 2024:
            for run in p.runs:
                run.text = run.text.replace("2024", str(an_referinta))

    doc.save(out_path)
    return out_path


# -------------------- CLI smoke test --------------------

def _smoke():
    sys.path.insert(0, str(Path(__file__).parent))
    from _paths import templates_dir
    tpl = templates_dir() / "Anexa-4-Calcul-intreprinderi.docx"
    out = Path.cwd() / f"_smoke_Anexa4_{date.today().isoformat()}.docx"
    parteneri = [
        {
            "denumire": "ALPHA SRL",
            "adresa": "Str. Test 1, București",
            "cui": "RO12345678",
            "reprezentant": "POPESCU ION",
            "procent": 30.0,
            "salariati": 25,
            "cifra_afaceri_mii_lei": 4500.123,
            "active_totale_mii_lei": 3200.456,
            "salariati_proportional": 7.5,
            "cifra_proportional": 1350.037,
            "active_proportional": 960.137,
        },
    ]
    legate = [
        {
            "denumire": "GAMMA SRL",
            "adresa": "Str. Demo 5, Cluj",
            "cui": "RO87654321",
            "reprezentant": "IONESCU MARIA",
            "salariati": 15,
            "cifra_afaceri_mii_lei": 2100.0,
            "active_totale_mii_lei": 1500.0,
        },
    ]
    agregat = {
        "A_salariati_sum": 7.5, "A_cifra_sum": 1350.037, "A_active_sum": 960.137,
        "B_salariati_sum": 15, "B_cifra_sum": 2100.0, "B_active_sum": 1500.0,
    }
    result = fill_anexa4(tpl, out, parteneri, legate, agregat, an_referinta=2024)
    print(f"Smoke OK: {result}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--smoke", action="store_true")
    args = ap.parse_args()
    if args.smoke:
        _smoke()
